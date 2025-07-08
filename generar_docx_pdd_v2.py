# -*- coding: utf-8 -*-
"""
Generador DOCX PDD v2.0 - Completamente Reestructurado
Implementa todos los requerimientos de mejora:
- Enfoque neutral As-Is (sin automatización)
- Eliminación de secciones To-Be 
- Tabla de contenidos automática
- Idioma consistente configurable
- Nueva estructura JSON (compatible con prompt v2.0)
- Reestructuración de pasos (tabla resumen + evidencia paso a paso)
- Nuevas secciones: stakeholders, métricas, dependencias
"""

import json
import os
import sys
import cv2
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime

# --- Configuración ---
JSON_INPUT_PATH = 'full_analysis_output.json'
SCREENSHOT_DIR = 'screenshots_output'
OUTPUT_DOCX_PATH = 'PDD_Generated_Output_v2.0.docx'

DEFAULT_USER_METADATA = {
    "project_name": "Business Process Documentation", 
    "project_acronym": "BPD",
    "author_name": "PDD Agent", 
    "version": "2.0", 
    "status": "DRAFT"
}

# --- Configuración de Idioma ---
LANGUAGE_CONFIG = {
    "es": {
        "document_title": "Documento de Descripción de Proceso (PDD)",
        "table_of_contents": "Tabla de Contenidos",
        "introduction": "Introducción",
        "purpose": "Propósito del Documento", 
        "business_context": "Contexto del Negocio",
        "process_description": "Descripción del Proceso",
        "overview": "Resumen General",
        "stakeholders": "Partes Interesadas",
        "applications": "Aplicaciones y Sistemas",
        "process_steps": "Pasos del Proceso",
        "step_summary": "Resumen de Pasos",
        "step_evidence": "Evidencia Paso a Paso",
        "inputs_outputs": "Entradas y Salidas",
        "business_rules": "Reglas de Negocio",
        "exceptions": "Excepciones Observadas",
        "metrics": "Métricas del Proceso",
        "dependencies": "Dependencias Técnicas",
        "quality_indicators": "Indicadores de Calidad",
        "appendix": "Apéndice",
        "glossary": "Glosario",
        "revision_history": "Historial de Revisiones",
        "references": "Referencias",
        "observed_only": "Se documentan únicamente elementos observados en el video.",
        "manual_completion": "Esta sección requiere completado manual por parte del analista de procesos.",
        "step_id": "ID",
        "application": "Aplicación", 
        "action": "Acción",
        "description": "Descripción Detallada",
        "screenshot": "Captura"
    },
    "en": {
        "document_title": "Process Description Document (PDD)",
        "table_of_contents": "Table of Contents", 
        "introduction": "Introduction",
        "purpose": "Document Purpose",
        "business_context": "Business Context", 
        "process_description": "Process Description",
        "overview": "Overview",
        "stakeholders": "Stakeholders",
        "applications": "Applications and Systems",
        "process_steps": "Process Steps",
        "step_summary": "Step Summary",
        "step_evidence": "Step-by-Step Evidence", 
        "inputs_outputs": "Inputs and Outputs",
        "business_rules": "Business Rules",
        "exceptions": "Observed Exceptions",
        "metrics": "Process Metrics",
        "dependencies": "Technical Dependencies",
        "quality_indicators": "Quality Indicators",
        "appendix": "Appendix",
        "glossary": "Glossary",
        "revision_history": "Revision History", 
        "references": "References",
        "observed_only": "Only elements observed in the video are documented.",
        "manual_completion": "This section requires manual completion by the process analyst.",
        "step_id": "ID",
        "application": "Application",
        "action": "Action", 
        "description": "Detailed Description",
        "screenshot": "Screenshot"
    }
}

# --- Funciones Auxiliares ---
def set_run_font(run, size_pt=11, bold=False, italic=False, color=None):
    """Aplica formato básico a un Run."""
    font = run.font
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = color

def add_manual_note(paragraph, lang_config):
    """Añade nota para secciones que requieren completado manual."""
    run_note = paragraph.add_run(f"[{lang_config['manual_completion']}]")
    set_run_font(run_note, 10, italic=True, color=RGBColor(128, 128, 128))

def setup_heading_styles(document):
    """Configura estilos de encabezado para ToC automática."""
    styles = document.styles
    
    # Asegurar que existan los estilos de título
    for i in range(1, 6):
        style_name = f'Heading {i}'
        if style_name not in styles:
            heading_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            heading_style.base_style = styles['Normal']
            
            # Configurar formato según nivel
            font = heading_style.font
            font.bold = True
            if i == 1:
                font.size = Pt(16)
                font.color.rgb = RGBColor(31, 56, 100)
            elif i == 2:
                font.size = Pt(14)
                font.color.rgb = RGBColor(47, 84, 150)
            elif i == 3:
                font.size = Pt(12)
                font.color.rgb = RGBColor(63, 112, 200)
            else:
                font.size = Pt(11)

def add_heading(document, text, level):
    """Añade encabezado con estilo apropiado para ToC."""
    heading = document.add_heading(text, level=level)
    return heading

def insert_table_of_contents(document, lang_config):
    """Inserta tabla de contenidos manual bien estructurada."""
    toc_heading = add_heading(document, lang_config["table_of_contents"], 1)
    
    # Crear tabla de contenidos manual
    toc_items = [
        ("1.", lang_config['introduction'], [
            ("1.1", lang_config['purpose']),
            ("1.2", lang_config['stakeholders']),
            ("1.3", lang_config['applications'])
        ]),
        ("2.", lang_config['business_context'], []),
        ("3.", lang_config['process_description'], [
            ("3.1", lang_config['step_summary']),
            ("3.2", lang_config['step_evidence'])
        ]),
        ("4.", "Análisis del Proceso", [
            ("4.1", lang_config['metrics']),
            ("4.2", lang_config['quality_indicators'])
        ]),
        ("5.", lang_config['inputs_outputs'], [
            ("5.1", "Entradas del Proceso"),
            ("5.2", "Salidas del Proceso")
        ]),
        ("6.", lang_config['business_rules'], []),
        ("7.", lang_config['exceptions'], []),
        ("8.", lang_config['dependencies'], [])
    ]
    
    # Crear tabla de contenidos visualmente atractiva
    for item_num, item_title, subitems in toc_items:
        # Título principal
        toc_para = document.add_paragraph()
        run_num = toc_para.add_run(f"{item_num} ")
        set_run_font(run_num, 12, bold=True, color=RGBColor(31, 56, 100))
        run_title = toc_para.add_run(item_title)
        set_run_font(run_title, 12, bold=True)
        
        # Subelementos
        for sub_num, sub_title in subitems:
            sub_para = document.add_paragraph()
            sub_para.paragraph_format.left_indent = Inches(0.5)
            run_sub_num = sub_para.add_run(f"{sub_num} ")
            set_run_font(run_sub_num, 11, color=RGBColor(47, 84, 150))
            run_sub_title = sub_para.add_run(sub_title)
            set_run_font(run_sub_title, 11)
    
    # Agregar nota informativa
    document.add_paragraph()
    note_para = document.add_paragraph()
    note_run = note_para.add_run("Nota: Esta tabla de contenidos refleja la estructura completa del documento generado automáticamente.")
    set_run_font(note_run, 9, italic=True, color=RGBColor(128, 128, 128))
    
    document.add_page_break()

# --- Funciones de Sección ---
def add_title_page(document, user_metadata, process_metadata, lang_config):
    """Añade página de título profesional."""
    print("  - Añadiendo página de título...")
    
    # Título principal
    title = document.add_heading(lang_config["document_title"], 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Nombre del proceso
    process_name = process_metadata.get('process_name', user_metadata.get('project_name', 'Unknown Process'))
    process_para = document.add_paragraph()
    process_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = process_para.add_run(process_name)
    set_run_font(run, 16, bold=True)
    
    # Metadata del documento
    document.add_paragraph()
    
    # Tabla de metadatos
    metadata_table = document.add_table(rows=5, cols=2)
    metadata_table.style = 'Table Grid'
    
    metadata_items = [
        ("Proyecto", user_metadata.get('project_name', 'N/A')),
        ("Autor", user_metadata.get('author_name', 'N/A')),
        ("Versión", user_metadata.get('version', 'N/A')),
        ("Estado", user_metadata.get('status', 'N/A')),
        ("Fecha", datetime.now().strftime("%Y-%m-%d"))
    ]
    
    for i, (label, value) in enumerate(metadata_items):
        metadata_table.cell(i, 0).text = label
        metadata_table.cell(i, 1).text = value
        set_run_font(metadata_table.cell(i, 0).paragraphs[0].runs[0], bold=True)
    
    document.add_page_break()

def add_stakeholders_section(document, stakeholders_data, lang_config):
    """Añade sección de partes interesadas/stakeholders."""
    print("  - Añadiendo sección de stakeholders...")
    
    if stakeholders_data:
        # Crear tabla de stakeholders
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Encabezados
        headers = ["Rol", "Responsabilidades", "Nivel de Evidencia"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            set_run_font(cell.paragraphs[0].runs[0], bold=True)
        
        # Datos
        for stakeholder in stakeholders_data:
            row = table.add_row()
            row.cells[0].text = stakeholder.get('role', 'N/A')
            row.cells[1].text = stakeholder.get('responsibilities', 'N/A')
            row.cells[2].text = stakeholder.get('evidence_level', 'N/A')
    else:
        p = document.add_paragraph()
        add_manual_note(p, lang_config)
        p.add_run(" No se identificaron stakeholders en el video.")
    
    document.add_paragraph()

def add_applications_section(document, applications_data, lang_config):
    """Añade sección de aplicaciones y dependencias."""
    print("  - Añadiendo sección de aplicaciones...")
    
    if applications_data:
        # Crear tabla de aplicaciones
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Encabezados
        headers = ["Aplicación", "Tipo", "Versión", "Crítica"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            set_run_font(cell.paragraphs[0].runs[0], bold=True)
        
        # Datos
        for app in applications_data:
            row = table.add_row()
            row.cells[0].text = app.get('application_name', 'N/A')
            row.cells[1].text = app.get('application_type', 'N/A')
            row.cells[2].text = app.get('version_visible', 'N/A') or 'N/A'
            row.cells[3].text = "Sí" if app.get('critical_for_process') else "No"
    else:
        p = document.add_paragraph()
        add_manual_note(p, lang_config)
        p.add_run(" No se identificaron aplicaciones.")
    
    document.add_paragraph()

def add_process_steps_summary_table(document, steps_data, lang_config):
    """Añade tabla resumen de pasos (sin screenshots)."""
    print("  - Añadiendo tabla resumen de pasos...")
    
    if not steps_data:
        p = document.add_paragraph()
        add_manual_note(p, lang_config)
        p.add_run(" No se encontraron pasos en el análisis.")
        return
    
    # Crear tabla resumen
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.autofit = False
    
    # Encabezados
    headers = [lang_config["step_id"], lang_config["application"], 
               lang_config["action"], lang_config["description"]]
    
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        set_run_font(cell.paragraphs[0].runs[0], bold=True)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Datos de pasos
    for step in steps_data:
        row = table.add_row()
        row.cells[0].text = str(step.get('step_number', 'N/A'))
        row.cells[1].text = step.get('application_in_focus', 'N/A')
        row.cells[2].text = step.get('action_summary', 'N/A')
        row.cells[3].text = step.get('detailed_description', 'N/A')
        
        # Formato de texto
        for cell in row.cells:
            if cell.text:
                set_run_font(cell.paragraphs[0].runs[0], 10)
    
    document.add_paragraph()

def add_step_by_step_evidence(document, steps_data, screenshot_dir, lang_config):
    """Añade sección de evidencia paso a paso con screenshots grandes."""
    print("  - Añadiendo evidencia paso a paso...")
    
    p = document.add_paragraph()
    p.add_run(f"{lang_config['observed_only']}")
    
    if not steps_data:
        add_manual_note(p, lang_config)
        return
    
    screenshots_found = 0
    
    for step in steps_data:
        step_num = step.get('step_number', 'N/A')
        action = step.get('action_summary', 'N/A')
        
        # Encabezado del paso
        step_heading = add_heading(document, f"Paso {step_num}: {action}", 3)
        
        # Descripción detallada
        desc_para = document.add_paragraph()
        desc_para.add_run("Descripción: ").bold = True
        desc_para.add_run(step.get('detailed_description', 'N/A'))
        
        # Detalles técnicos
        if step.get('ui_elements_details'):
            ui_para = document.add_paragraph()
            ui_para.add_run("Elementos UI: ").bold = True
            ui_para.add_run(step.get('ui_elements_details', 'N/A'))
        
        if step.get('data_manipulated'):
            data_para = document.add_paragraph()
            data_para.add_run("Datos manipulados: ").bold = True 
            data_para.add_run(step.get('data_manipulated', 'N/A'))
        
        # Screenshot
        screenshot_filename = f"screenshot_paso_{step_num}.png"
        screenshot_path = os.path.join(screenshot_dir, screenshot_filename)
        
        if os.path.exists(screenshot_path):
            try:
                img_para = document.add_paragraph()
                img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = img_para.add_run()
                run.add_picture(screenshot_path, width=Inches(6))  # Imagen grande y legible
                screenshots_found += 1
                
                # Caption
                caption_para = document.add_paragraph()
                caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                caption_run = caption_para.add_run(f"Figura {step_num}: {action}")
                set_run_font(caption_run, 9, italic=True)
                
            except Exception as e:
                print(f"    Error al insertar screenshot {screenshot_path}: {e}")
                error_para = document.add_paragraph()
                error_para.add_run(f"[Error al cargar screenshot: {e}]")
                set_run_font(error_para.runs[0], 9, italic=True, color=RGBColor(255, 0, 0))
        else:
            no_img_para = document.add_paragraph()
            no_img_para.add_run("[Screenshot no disponible]")
            set_run_font(no_img_para.runs[0], 9, italic=True, color=RGBColor(128, 128, 128))
        
        document.add_paragraph()  # Espaciado
    
    print(f"    -> Evidencia creada para {len(steps_data)} pasos")
    print(f"    -> {screenshots_found} screenshots insertados")

def add_metrics_section(document, metrics_data, lang_config):
    """Añade sección de métricas del proceso."""
    print("  - Añadiendo métricas del proceso...")
    
    if metrics_data:
        # Crear tabla de métricas
        metrics_list = [
            ("Tiempo total de ejecución", f"{metrics_data.get('total_execution_time_seconds', 'N/A')} segundos"),
            ("Número de pasos manuales", str(metrics_data.get('manual_steps_count', 'N/A'))),
            ("Interacciones con sistema", str(metrics_data.get('system_interactions_count', 'N/A'))),
            ("Pasos de entrada de datos", str(metrics_data.get('data_entry_steps_count', 'N/A'))),
            ("Pasos de validación", str(metrics_data.get('validation_steps_count', 'N/A')))
        ]
        
        table = document.add_table(rows=len(metrics_list), cols=2)
        table.style = 'Table Grid'
        
        for i, (metric, value) in enumerate(metrics_list):
            table.cell(i, 0).text = metric
            table.cell(i, 1).text = value
            set_run_font(table.cell(i, 0).paragraphs[0].runs[0], bold=True)
    else:
        p = document.add_paragraph()
        add_manual_note(p, lang_config)
        p.add_run(" No se generaron métricas automáticas.")
    
    document.add_paragraph()

def add_quality_indicators_section(document, quality_data, lang_config):
    """Añade sección de indicadores de calidad.""" 
    print("  - Añadiendo indicadores de calidad...")
    
    if quality_data:
        indicators = [
            ("Dudas del usuario observadas", "Sí" if quality_data.get('user_hesitation_observed') else "No"),
            ("Número de correcciones de errores", str(quality_data.get('error_corrections_count', 0))),
            ("Acciones repetidas", str(quality_data.get('repeated_actions_count', 0))),
            ("Búsqueda de ayuda", "Sí" if quality_data.get('help_seeking_behavior') else "No")
        ]
        
        table = document.add_table(rows=len(indicators), cols=2)
        table.style = 'Table Grid'
        
        for i, (indicator, value) in enumerate(indicators):
            table.cell(i, 0).text = indicator
            table.cell(i, 1).text = value
            set_run_font(table.cell(i, 0).paragraphs[0].runs[0], bold=True)
    else:
        p = document.add_paragraph()
        add_manual_note(p, lang_config)
        p.add_run(" No se generaron indicadores de calidad.")
    
    document.add_paragraph()

def generate_pdd_docx_v2(json_path: str, screenshot_dir: str, output_docx_path: str, 
                        user_metadata: dict, language: str = "es"):
    """
    Genera el documento DOCX del PDD v2.0 usando el JSON de análisis.
    
    Args:
        json_path: Ruta al archivo JSON con el análisis
        screenshot_dir: Directorio con las capturas de pantalla
        output_docx_path: Ruta de salida del documento DOCX
        user_metadata: Metadatos del usuario/proyecto
        language: Idioma del documento ("es" o "en")
    """
    
    print(f"\n--- Generando PDD DOCX v2.0 ({language.upper()}) ---")
    print(f"JSON Input: {json_path}")
    print(f"Screenshots: {screenshot_dir}")
    print(f"Output: {output_docx_path}")
    
    # Validar archivos de entrada
    if not os.path.exists(json_path):
        print(f"❌ ERROR: Archivo JSON no encontrado: {json_path}")
        return False
        
    # Cargar configuración de idioma
    lang_config = LANGUAGE_CONFIG.get(language, LANGUAGE_CONFIG["es"])
    
    # Cargar JSON de análisis
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            analysis_data = json.load(f)
        print(f"✅ JSON cargado exitosamente")
    except Exception as e:
        print(f"❌ ERROR al cargar JSON: {e}")
        return False
    
    # Crear documento
    try:
        document = Document()
        setup_heading_styles(document)
        
        # === PÁGINA DE TÍTULO ===
        add_title_page(document, user_metadata, analysis_data.get('pdd_metadata', {}), lang_config)
        
        # === TABLA DE CONTENIDOS ===
        insert_table_of_contents(document, lang_config)
        
        # === 1. INTRODUCCIÓN ===
        add_heading(document, f"1. {lang_config['introduction']}", 1)
        
        # 1.1 Propósito del Documento
        add_heading(document, f"1.1 {lang_config['purpose']}", 2)
        purpose_para = document.add_paragraph()
        purpose_para.add_run(
            f"Este documento describe el proceso de negocio '{analysis_data.get('pdd_metadata', {}).get('process_name', 'Proceso No Definido')}' "
            f"tal como se ejecuta actualmente (As-Is). El objetivo es proporcionar una documentación precisa "
            f"del estado actual del proceso para facilitar su comprensión, capacitación, auditoría o análisis de mejora."
        )
        
        # 1.2 Stakeholders
        add_heading(document, f"1.2 {lang_config['stakeholders']}", 2)
        add_stakeholders_section(document, analysis_data.get('stakeholders_identified', []), lang_config)
        
        # 1.3 Aplicaciones
        add_heading(document, f"1.3 {lang_config['applications']}", 2)
        add_applications_section(document, analysis_data.get('applications_dependencies', []), lang_config)
        
        # === 2. CONTEXTO DEL NEGOCIO ===
        add_heading(document, f"2. {lang_config['business_context']}", 1)
        context_para = document.add_paragraph()
        process_context = analysis_data.get('process_context', {})
        context_para.add_run(f"Propósito: {process_context.get('business_purpose', 'No especificado')}")
        context_para.add_run(f"\nÁrea/Departamento: {process_context.get('department_area', 'No especificado')}")
        context_para.add_run(f"\nPropietario del Proceso: {process_context.get('process_owner_role', 'No especificado')}")
        
        # === 3. DESCRIPCIÓN DEL PROCESO ===
        add_heading(document, f"3. {lang_config['process_description']}", 1)
        
        # 3.1 Tabla resumen de pasos
        add_heading(document, f"3.1 {lang_config['step_summary']}", 2)
        add_process_steps_summary_table(document, analysis_data.get('process_steps', []), lang_config)
        
        # 3.2 Evidencia paso a paso
        add_heading(document, f"3.2 {lang_config['step_evidence']}", 2)
        add_step_by_step_evidence(document, analysis_data.get('process_steps', []), screenshot_dir, lang_config)
        
        # === 4. ANÁLISIS DEL PROCESO ===
        add_heading(document, "4. Análisis del Proceso", 1)
        
        # 4.1 Métricas
        add_heading(document, f"4.1 {lang_config['metrics']}", 2)
        add_metrics_section(document, analysis_data.get('process_metrics', {}), lang_config)
        
        # 4.2 Indicadores de calidad
        add_heading(document, f"4.2 {lang_config['quality_indicators']}", 2)
        add_quality_indicators_section(document, analysis_data.get('quality_indicators', {}), lang_config)
        
        # === 5. ENTRADAS Y SALIDAS ===
        add_heading(document, f"5. {lang_config['inputs_outputs']}", 1)
        
        # Entradas
        add_heading(document, "5.1 Entradas del Proceso", 2)
        inputs_para = document.add_paragraph()
        inputs_data = analysis_data.get('process_inputs', [])
        if inputs_data:
            for inp in inputs_data:
                inputs_para.add_run(f"• {inp.get('input_name', 'N/A')} ({inp.get('input_type', 'N/A')})\n")
        else:
            add_manual_note(inputs_para, lang_config)
        
        # Salidas
        add_heading(document, "5.2 Salidas del Proceso", 2)
        outputs_para = document.add_paragraph()
        outputs_data = analysis_data.get('process_outputs', [])
        if outputs_data:
            for out in outputs_data:
                outputs_para.add_run(f"• {out.get('output_name', 'N/A')} ({out.get('output_type', 'N/A')})\n")
        else:
            add_manual_note(outputs_para, lang_config)
        
        # === 6. REGLAS DE NEGOCIO ===
        add_heading(document, f"6. {lang_config['business_rules']}", 1)
        rules_para = document.add_paragraph()
        rules_data = analysis_data.get('business_rules_observed', [])
        if rules_data:
            for rule in rules_data:
                rules_para.add_run(f"• {rule.get('rule_description', 'N/A')}\n")
        else:
            add_manual_note(rules_para, lang_config)
        
        # === 7. EXCEPCIONES ===
        add_heading(document, f"7. {lang_config['exceptions']}", 1)
        exceptions_para = document.add_paragraph()
        exceptions_data = analysis_data.get('exceptions_observed', [])
        if exceptions_data:
            for exc in exceptions_data:
                exceptions_para.add_run(f"• {exc.get('description', 'N/A')} - {exc.get('handling_method', 'N/A')}\n")
        else:
            exceptions_para.add_run("No se observaron excepciones durante la ejecución del proceso.")
        
        # === 8. DEPENDENCIAS TÉCNICAS ===
        add_heading(document, f"8. {lang_config['dependencies']}", 1)
        deps_para = document.add_paragraph()
        deps_data = analysis_data.get('technical_dependencies', [])
        if deps_data:
            for dep in deps_data:
                criticality = dep.get('criticality', 'N/A')
                deps_para.add_run(f"• {dep.get('dependency_name', 'N/A')} ({dep.get('dependency_type', 'N/A')}) - Criticidad: {criticality}\n")
        else:
            add_manual_note(deps_para, lang_config)
        
        # Guardar documento
        document.save(output_docx_path)
        
        print(f"✅ Documento PDD v2.0 generado exitosamente: {output_docx_path}")
        
        # Estadísticas finales
        file_size_mb = os.path.getsize(output_docx_path) / (1024 * 1024)
        print(f"📊 Tamaño del archivo: {file_size_mb:.2f} MB")
        
        step_count = len(analysis_data.get('process_steps', []))
        print(f"📊 Pasos documentados: {step_count}")
        
        return True
        
    except Exception as e:
        print(f"❌ ERROR al generar documento: {e}")
        return False

# === FUNCIÓN PRINCIPAL ===
def main():
    """Función principal para pruebas."""
    
    # Metadatos del usuario/proyecto  
    user_metadata = {
        "project_name": "Proceso de Ejemplo PDD v2.0",
        "project_acronym": "PDD",
        "author_name": "Analista PDD Agent",
        "version": "2.0",
        "status": "DRAFT"
    }
    
    # Generar documento
    success = generate_pdd_docx_v2(
        json_path=JSON_INPUT_PATH,
        screenshot_dir=SCREENSHOT_DIR, 
        output_docx_path=OUTPUT_DOCX_PATH,
        user_metadata=user_metadata,
        language="es"  # Cambiar a "en" para inglés
    )
    
    if success:
        print(f"\n🎉 Generación completada exitosamente!")
        print(f"📄 Documento: {OUTPUT_DOCX_PATH}")
    else:
        print(f"\n❌ Error en la generación del documento")

if __name__ == "__main__":
    main() 